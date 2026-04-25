from __future__ import annotations

from datetime import datetime, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from io import BytesIO
import json
import os
import random
import smtplib
import string
import uuid

from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, Response, StreamingResponse

try:
    from openai import OpenAI
except Exception:  # pragma: no cover
    OpenAI = None

try:
    from anthropic import Anthropic
except Exception:  # pragma: no cover
    Anthropic = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except Exception:  # pragma: no cover
    Document = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    qn = None
    Pt = None
    RGBColor = None


APP_VERSION = "1.7.16"

PRIMARY_EMAIL = os.environ.get("ADMIN_EMAIL", "info@boldkimya.com.tr")
GMAIL_USER = os.environ.get("GMAIL_USER", "")
GMAIL_PASS = os.environ.get("GMAIL_APP_PASSWORD", "")

PASSWORD = "Q7m!R2x#"

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "").strip()
OPENAI_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini").strip()

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "").strip()
ANTHROPIC_MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-5").strip()

XAI_API_KEY = os.environ.get("XAI_API_KEY", "").strip()
XAI_MODEL = os.environ.get("XAI_MODEL", "grok-4-1-fast-non-reasoning").strip()
XAI_BASE_URL = os.environ.get("XAI_BASE_URL", "https://api.x.ai/v1").strip()

LLAMA_BASE_URL = os.environ.get("LLAMA_BASE_URL", "").strip()
LLAMA_MODEL = os.environ.get("LLAMA_MODEL", "").strip()
LLAMA_API_KEY = os.environ.get("LLAMA_API_KEY", "local-dev-key").strip() or "local-dev-key"

app = FastAPI(title="T.C. ANATOLIA-Q", version=APP_VERSION)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

USERS = {
    "158963": {"name": "Sistem Yonetici", "role": "admin"},
    "274851": {"name": "Operasyon Birimi", "role": "operator"},
    "386472": {"name": "Stratejik Analiz", "role": "analyst"},
    "491205": {"name": "Enerji Izleme", "role": "analyst"},
    "563184": {"name": "Saha Operatoru", "role": "operator"},
}

DOMAINS = {
    "savunma": {"display": "Savunma", "kurumlar": ["MSB", "TSK", "MIT"]},
    "ekonomi": {"display": "Ekonomi", "kurumlar": ["Hazine ve Maliye Bakanlığı", "TCMB", "BDDK"]},
    "enerji": {"display": "Enerji", "kurumlar": ["Enerji ve Tabii Kaynaklar Bakanlığı", "EPDK", "BOTAŞ"]},
    "dis_politika": {"display": "Dış Politika", "kurumlar": ["Dışişleri Bakanlığı", "Cumhurbaşkanlığı", "MIT"]},
    "toplumsal_olaylar": {"display": "Toplumsal Olaylar", "kurumlar": ["İçişleri Bakanlığı", "Emniyet Genel Müdürlüğü", "Valilikler"]},
    "genel_chat": {"display": "Genel Chat", "kurumlar": ["Cumhurbaşkanlığı", "Strateji Birimi", "Merkez Koordinasyon"]},
    "cross": {"display": "Çapraz Sentez", "kurumlar": ["Cumhurbaşkanlığı", "MSB", "Dışişleri Bakanlığı"]},
}

PROVIDER_ROUTE = {
    "genel_chat": ["anthropic", "openai", "grok", "llama", "fallback"],
    "savunma": ["anthropic", "openai", "grok", "llama", "fallback"],
    "ekonomi": ["anthropic", "openai", "grok", "llama", "fallback"],
    "enerji": ["anthropic", "openai", "grok", "llama", "fallback"],
    "dis_politika": ["anthropic", "openai", "grok", "llama", "fallback"],
    "toplumsal_olaylar": ["anthropic", "openai", "grok", "llama", "fallback"],
    "cross": ["anthropic", "openai", "grok", "llama", "fallback"],
}

pending_codes = {}
active_sessions = {}
analysis_store = {}
alerts_store = []
ops_feed_store = []
provider_cooldowns = {}

TEXT_REPLACEMENTS = [
    ("Ä±", "ı"),
    ("Ä°", "İ"),
    ("ÄŸ", "ğ"),
    ("Äž", "Ğ"),
    ("ÅŸ", "ş"),
    ("Åž", "Ş"),
    ("Ã¼", "ü"),
    ("Ãœ", "Ü"),
    ("Ã¶", "ö"),
    ("Ã–", "Ö"),
    ("Ã§", "ç"),
    ("Ã‡", "Ç"),
    ("â€™", "'"),
    ("â€œ", '"'),
    ("â€", '"'),
    ("â€¦", "..."),
]


def stamp():
    return datetime.now().strftime("%d.%m.%Y %H:%M")


def send_mail(subject: str, html: str):
    if not GMAIL_USER or not GMAIL_PASS:
        raise HTTPException(500, "E-posta ayarlar eksik.")

    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"] = GMAIL_USER
    msg["To"] = PRIMARY_EMAIL
    msg.attach(MIMEText(html, "html", "utf-8"))

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(GMAIL_USER, GMAIL_PASS)
        smtp.sendmail(GMAIL_USER, [PRIMARY_EMAIL], msg.as_string())


def safe_send_mail(subject: str, html: str):
    try:
        send_mail(subject, html)
    except HTTPException:
        return False
    return True


def token_from(request: Request, body=None):
    auth = request.headers.get("authorization", "")
    if auth.lower().startswith("bearer "):
        return auth.split(" ", 1)[1].strip()
    return request.headers.get("x-auth-token", "").strip() or str((body or {}).get("token", "")).strip()


def require_session(request: Request, body=None):
    session = active_sessions.get(token_from(request, body))
    if not session:
        raise HTTPException(401, "Gecersiz oturum.")
    return session


def clean_name(value):
    raw = "".join(ch for ch in str(value or "").strip() if ch.isalnum() or ch in " .-_")
    return raw[:24].strip() or "dostum"


def normalize_text(value):
    text = str(value or "")
    for old, new in TEXT_REPLACEMENTS:
        text = text.replace(old, new)
    return text


def normalize_payload(value):
    if isinstance(value, str):
        return normalize_text(value)
    if isinstance(value, list):
        return [normalize_payload(item) for item in value]
    if isinstance(value, dict):
        return {key: normalize_payload(item) for key, item in value.items()}
    return value


def trim_store(store, limit=60):
    while len(store) > limit:
        store.pop(0)


def extract_json_text(value: str) -> str:
    text = str(value or "").strip()
    if text.startswith("```"):
        lines = text.splitlines()
        lines = lines[1:]
        while lines and lines[-1].strip().startswith("```"):
            lines.pop()
        text = "\n".join(lines).strip()

    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        return text[start : end + 1]
    return text


def is_soft_provider_failure(error: Exception) -> bool:
    message = str(error).lower()
    quota_markers = [
        "quota",
        "rate limit",
        "insufficient_quota",
        "billing",
        "credit",
        "429",
        "overloaded",
        "capacity",
        "temporarily unavailable",
    ]
    return any(marker in message for marker in quota_markers)


def provider_available(name: str) -> bool:
    until = provider_cooldowns.get(name)
    return not until or datetime.now() >= until


def cool_down_provider(name: str, minutes: int = 10):
    provider_cooldowns[name] = datetime.now() + timedelta(minutes=minutes)


def provider_status():
    return {
        "openai": bool(OpenAI and OPENAI_API_KEY) and provider_available("openai"),
        "anthropic": bool(Anthropic and ANTHROPIC_API_KEY) and provider_available("anthropic"),
        "grok": bool(OpenAI and XAI_API_KEY and XAI_MODEL and XAI_BASE_URL) and provider_available("grok"),
        "llama": bool(OpenAI and LLAMA_BASE_URL and LLAMA_MODEL) and provider_available("llama"),
        "fallback": True,
    }


def build_schema_hint(domain: str) -> str:
    if domain == "genel_chat":
        return json.dumps(
            {
                "ozet": "Ksa ve dogal sohbet cevab",
                "tehdit_analizi": "Konuya dair ksa degerlendirme",
                "senaryolar": [
                    "Bunu daha sade anlat.",
                    "Bana 3 maddede ozetle.",
                    "Bir tk daha ciddi tonda yeniden yaz.",
                ],
                "oncelikli_oneri": "Bir sonraki mesaj icin oneri",
                "etkilenen_kurumlar": ["Genel Bilgi", "Gundelik Dil", "Hzl Ozet"],
                "zaman_cercevesi": "Anlk sohbet",
                "sohbet_tonu": "Rahat, akc ve hafif sakac",
                "kritik_baglanti": "Konunun devam baglants",
                "tehdit_seviyesi": "DUSUK",
            },
            ensure_ascii=False,
            indent=2,
        )

    payload = {
        "ozet": "Yonetici ozeti",
        "tehdit_analizi": "Durumun stratejik analizi",
        "senaryolar": [
            {
                "baslik": "Senaryo baslg",
                "olasilik": "Yuksek",
                "aciklama": "Senaryonun ksa acklamas",
                "aksiyon": "Onerilen aksiyon",
            }
        ],
        "oncelikli_oneri": "Birincil oneri",
        "etkilenen_kurumlar": ["Kurum 1", "Kurum 2"],
        "zaman_cercevesi": "Acil / Ksa / Orta / Uzun",
        "kritik_baglanti": "Gozlenmesi gereken kilit bag",
        "tehdit_seviyesi": "DUSUK|ORTA|YUKSEK|KRITIK",
    }

    if domain == "cross":
        payload["genel_tehdit_seviyesi"] = "ORTA"
        payload["alan_etkileri"] = {
            "savunma": {"etki": "orta", "aciklama": "Alan etkisi"},
            "ekonomi": {"etki": "orta", "aciklama": "Alan etkisi"},
            "enerji": {"etki": "orta", "aciklama": "Alan etkisi"},
            "toplumsal_olaylar": {"etki": "orta", "aciklama": "Alan etkisi"},
            "dis_politika": {"etki": "orta", "aciklama": "Alan etkisi"},
        }
    return json.dumps(payload, ensure_ascii=False, indent=2)


def build_system_prompt(domain: str) -> str:
    area = DOMAINS[domain]["display"]
    return (
        "Sen T.C. ANATOLIA-Q icin calsan resmi karar destek motorusun. "
        "Yantlarn Turkce olmal, duz uydurma yapmamal, risk ve oneri ayrmn net tutmal. "
        "Kullanc arayuzu sadece JSON bekliyor; asla acklama, markdown, kod blogu veya onsoz ekleme. "
        f"Alan: {area}. "
        "Resmi kurum dili kullan ama okunabilir ve net kal. "
        "Belirsiz bilgi varsa bunu tahmini degil dikkatli bir dille isle. "
        "Ckt yaps su ornegi izlemeli:\n"
        f"{build_schema_hint(domain)}"
    )


def build_user_prompt(domain: str, situation: str, chat_name: str = "") -> str:
    if domain == "genel_chat":
        name = clean_name(chat_name)
        return (
            f"Hitap ad: {name}\n"
            "Bu mod sohbet ekran icindir. Cevap dogal, ksa-orta uzunlukta, akc ve hafif sakac olsun. "
            "Yine de sacmalama, soruyu gercekten cevapla.\n"
            f"Kullanc mesaj:\n{situation.strip()}"
        )

    area = DOMAINS[domain]["display"]
    institutions = ", ".join(DOMAINS[domain]["kurumlar"])
    return (
        f"Analiz alan: {area}\n"
        f"Ilgili kurumlar: {institutions}\n"
        "Asagdaki durumu karar destek mantgyla analiz et. "
        "Ksa vadeli risk, senaryo ve kurumlar aras etkileri ayr.\n"
        f"Durum notu:\n{situation.strip()}"
    )


def normalize_general_chat(raw: dict, chat_name: str = "") -> dict:
    name = clean_name(chat_name)
    scenarios = raw.get("senaryolar") or [
        "Bunu daha sade anlat.",
        "Bana 3 maddede ozetle.",
        "Bir tk daha ciddi tonda yeniden yaz.",
    ]
    if not isinstance(scenarios, list):
        scenarios = [str(scenarios)]

    return {
        "ozet": str(raw.get("ozet") or f"Selam {name}, devam edelim. Sorunu biraz daha acarsan daha net gidebilirim."),
        "tehdit_analizi": str(raw.get("tehdit_analizi") or "Konuya dair kisa degerlendirme hazirlandi."),
        "senaryolar": [str(item) for item in scenarios[:3]],
        "oncelikli_oneri": str(raw.get("oncelikli_oneri") or "Bir sonraki mesajda tek soru sor; cevab keskinlestireyim."),
        "etkilenen_kurumlar": raw.get("etkilenen_kurumlar") or ["Genel Bilgi", "Gundelik Dil", "Hzl Ozet"],
        "zaman_cercevesi": str(raw.get("zaman_cercevesi") or "Anlk sohbet"),
        "sohbet_tonu": str(raw.get("sohbet_tonu") or "Rahat, akici ve hafif sakaci"),
        "kritik_baglanti": str(raw.get("kritik_baglanti") or "Istersen ayn konuda devam edebiliriz."),
        "tehdit_seviyesi": str(raw.get("tehdit_seviyesi") or "DUSUK").upper(),
    }


def normalize_analysis(domain: str, raw: dict, chat_name: str = "") -> dict:
    if domain == "genel_chat":
        return normalize_general_chat(raw, chat_name)

    scenarios = raw.get("senaryolar") or []
    normalized_scenarios = []
    for item in scenarios[:3]:
        if isinstance(item, dict):
            normalized_scenarios.append(
                {
                    "baslik": str(item.get("baslik") or "Senaryo"),
                    "olasilik": str(item.get("olasilik") or "Orta").capitalize(),
                    "aciklama": str(item.get("aciklama") or "Ek degerlendirme gerekiyor."),
                    "aksiyon": str(item.get("aksiyon") or "Yakn izleme onerilir."),
                }
            )
        else:
            normalized_scenarios.append(
                {
                    "baslik": "Senaryo",
                    "olasilik": "Orta",
                    "aciklama": str(item),
                    "aksiyon": "Yakn izleme onerilir.",
                }
            )

    while len(normalized_scenarios) < 3:
        normalized_scenarios.append(
            {
                "baslik": f"Senaryo {len(normalized_scenarios) + 1}",
                "olasilik": "Orta",
                "aciklama": "Ek modelleme ile netlestirilmeli.",
                "aksiyon": "Veri takibi surdurulmeli.",
            }
        )

    result = {
        "ozet": str(raw.get("ozet") or f"{DOMAINS[domain]['display']} icin degerlendirme uretildi."),
        "tehdit_analizi": str(raw.get("tehdit_analizi") or "Durum analizi tamamland."),
        "senaryolar": normalized_scenarios,
        "oncelikli_oneri": str(raw.get("oncelikli_oneri") or "Kurumlar aras esgudum korunmal."),
        "etkilenen_kurumlar": raw.get("etkilenen_kurumlar") or DOMAINS[domain]["kurumlar"],
        "zaman_cercevesi": str(raw.get("zaman_cercevesi") or "Acil"),
        "kritik_baglanti": str(raw.get("kritik_baglanti") or "Alanlar aras baglar duzenli izlenmeli."),
        "tehdit_seviyesi": str(raw.get("tehdit_seviyesi") or "ORTA").upper(),
    }

    if domain == "cross":
        alan_etkileri = raw.get("alan_etkileri") or {}
        result["genel_tehdit_seviyesi"] = str(raw.get("genel_tehdit_seviyesi") or result["tehdit_seviyesi"]).upper()
        result["alan_etkileri"] = {
            key: {
                "etki": str((alan_etkileri.get(key) or {}).get("etki") or "orta"),
                "aciklama": str((alan_etkileri.get(key) or {}).get("aciklama") or "Takip edilmelidir."),
            }
            for key in ["savunma", "ekonomi", "enerji", "toplumsal_olaylar", "dis_politika"]
        }

    return result


def general_chat_reply(situation, chat_name=""):
    name = clean_name(chat_name)
    text = " ".join(str(situation or "").split())
    low = text.lower()

    if any(word in low for word in ["selam", "merhaba", "sa", "naber", "naslsn", "nasilsin"]):
        answer = f"Selam {name}, ne konusmak istiyorsun?"
    elif any(word in low for word in ["nedir", "ne demek", "anlat", "ackla", "acikla"]):
        answer = f"{name}, bunu sade anlatalm: {text[:220]}. Ksa cevap su; konu parcalar dogru yere oturtma isi."
    elif "?" in text or any(word in low for word in ["neden", "niye", "nasl", "nasil", "kim", "ne zaman", "hangi"]):
        answer = f"{name}, hzl cevap vereyim: once resmi gor, sonra parcala, sonra en guclu noktay sec."
    else:
        answer = f"{name}, notunu aldm. Bunu fazla kasmadan toparlayaym: {text[:220]}. Istersen tonu daha ciddi ya da daha eglenceli yapabilirim."

    return normalize_general_chat(
        {
            "ozet": answer,
            "tehdit_analizi": "Konuya dair kisa degerlendirme hazirlandi.",
            "senaryolar": [
                "Bunu daha sade anlat.",
                "Bana 3 maddede ozetle.",
                "Bir tk daha ciddi tonda yeniden yaz.",
            ],
            "oncelikli_oneri": "Bir sonraki mesajda tek bir soru ya da konu baslg at; cevab daha keskinlestireyim.",
            "etkilenen_kurumlar": ["Genel Bilgi", "Gundelik Dil", "Hzl Ozet"],
            "zaman_cercevesi": "Anlk sohbet",
            "sohbet_tonu": "Rahat, akici ve hafif sakaci",
            "kritik_baglanti": "Ayn konuyu daha ciddi, daha ksa ya da daha eglenceli tonda surdurebiliriz.",
            "tehdit_seviyesi": "DUSUK",
        },
        chat_name,
    )


def fallback(domain, situation, chat_name=""):
    if domain == "genel_chat":
        return general_chat_reply(situation, chat_name)

    base = DOMAINS[domain]
    result = {
        "ozet": f"{base['display']} icin degerlendirme uretildi. Ana eksen: {situation[:220]}",
        "tehdit_analizi": "Mevcut bulgular cercevesinde durum degerlendirmesi sunulmustur.",
        "senaryolar": [
            {
                "baslik": "Gerilim artar",
                "olasilik": "Yuksek",
                "aciklama": "Ksa vadede bask artabilir.",
                "aksiyon": "Anlk izleme ve koordinasyon surdurulmeli.",
            },
            {
                "baslik": "Etki dengelenir",
                "olasilik": "Orta",
                "aciklama": "Hzl tepki ile etki snrlanabilir.",
                "aksiyon": "Durum raporlamas sklastrlmal.",
            },
            {
                "baslik": "Etki daralr",
                "olasilik": "Dusuk",
                "aciklama": "Tetikleyiciler zayflarsa tablo yumusayabilir.",
                "aksiyon": "Yedek planlar hazr tutulmal.",
            },
        ],
        "oncelikli_oneri": "Kurumlar aras koordinasyon korunmal ve durum izlenmelidir.",
        "etkilenen_kurumlar": base["kurumlar"],
        "zaman_cercevesi": "Acil",
        "kritik_baglanti": "Guncel veri geldikce analiz derinlestirilmeli.",
        "tehdit_seviyesi": "ORTA",
    }

    if domain == "cross":
        result["genel_tehdit_seviyesi"] = "ORTA"
        result["alan_etkileri"] = {
            key: {"etki": "orta", "aciklama": "Takip edilmelidir."}
            for key in ["savunma", "ekonomi", "enerji", "toplumsal_olaylar", "dis_politika"]
        }

    return normalize_analysis(domain, result, chat_name)


def report_code_for(domain: str) -> str:
    codes = {
        "savunma": "SAV",
        "ekonomi": "EKO",
        "enerji": "ENR",
        "dis_politika": "DIP",
        "toplumsal_olaylar": "TOP",
        "genel_chat": "GCH",
        "cross": "SEN",
    }
    return codes.get(domain, "GEN")


def report_title_for(domain: str) -> str:
    titles = {
        "savunma": "ULUSAL SAVUNMA DURUM DEGERLENDIRMESI",
        "ekonomi": "EKONOMIK ETKI VE DAYANIKLILIK RAPORU",
        "enerji": "ENERJI GUVENLIGI VE SUREKLILIK RAPORU",
        "dis_politika": "DIS POLITIKA VE BOLGESEL ETKI RAPORU",
        "toplumsal_olaylar": "TOPLUMSAL OLAYLAR DURUM DEGERLENDIRMESI",
        "genel_chat": "GENEL CHAT GORUSME NOTU",
        "cross": "CAPRAZ ETKI STRATEJIK SENTEZ RAPORU",
    }
    return titles.get(domain, "STRATEJIK DEGERLENDIRME RAPORU")


def report_profile_for(domain: str) -> dict:
    profiles = {
        "savunma": {
            "cover_title": "ULUSAL SAVUNMA VE TEHDIT DEGERLENDIRME RAPORU",
            "scope": "Ulusal guvenlik, snr hatt, askeri hazrlk ve caydrclk gorunumu",
            "capacity": "Mevcut savunma gorunumu; saha sinyalleri, kurumlar aras esgudum seviyesi ve ksa vadeli reaksiyon kapasitesi uzerinden degerlendirilmistir.",
            "architecture": "Cok katmanl erken uyar, teyit edilmis saha beslemesi, musterek komuta aks ve hzl karar dongusu birlikte ele alnmaldr.",
            "standards": [
                "Kritik savunma verileri tek operasyon masasnda birlestirilmelidir.",
                "Uyar, teyit ve gorevlendirme ayn karar zinciri icinde tutulmaldr.",
                "Merkez bildirimleri kullanc kodu, zaman damgas ve oncelik derecesi ile kayt altna alnmaldr.",
            ],
        },
        "ekonomi": {
            "cover_title": "EKONOMIK ETKI, DAYANIKLILIK VE RISK RAPORU",
            "scope": "Makroekonomik gorunum, piyasa etkisi, finansal istikrar ve kurumsal tepki alan",
            "capacity": "Ekonomik gorunum; piyasa hassasiyetleri, kurumsal tamponlar ve karar alclarn manevra alan uzerinden ele alnmstr.",
            "architecture": "Erken uyari gostergeleri, sektor bazl sinyal takibi ve merkez teyit zinciri ayn tabloda tutulmaldr.",
            "standards": [
                "Kurumsal etki analizi finansal piyasa verisi ile birlikte degerlendirilmelidir.",
                "Kritik ekonomik sinyaller gunluk degil olay bazl takip mantgyla ele alnmaldr.",
                "Ekonomi, enerji ve toplumsal tepki baglants ayn rapor ekseninde tutulmaldr.",
            ],
        },
        "enerji": {
            "cover_title": "ENERJI GUVENLIGI VE SUREKLILIK RAPORU",
            "scope": "Arz guvenligi, iletim-uretim surekliligi ve kritik altyap dayanklg",
            "capacity": "Enerji gorunumu; tedarik zinciri, kritik saha bagmlklari ve kesinti tolerans uzerinden degerlendirilmistir.",
            "architecture": "Tedarik, iletim, depolama ve kriz yedekliligi tek operasyon gordesinda birlikte izlenmelidir.",
            "standards": [
                "Arz surekliligi uzerindeki basklar erken evrede alarm zincirine girmelidir.",
                "Enerji etkileri toplumsal ve ekonomik yansmalarla birlikte raporlanmaldr.",
                "Kritik altyap noktalar icin bolgesel oncelik haritas tutulmaldr.",
            ],
        },
        "dis_politika": {
            "cover_title": "DIS POLITIKA VE BOLGESEL ETKI RAPORU",
            "scope": "Bolgesel diplomasi, dis bask unsurlari ve stratejik ortaklik dengesi",
            "capacity": "Dis politika alan; devletler aras etki, diplomatik hareket alan ve ikincil guvenlik yansmalar uzerinden okunmustur.",
            "architecture": "Diplomatik sinyal, sahadaki etkiler ve kurumlar arasi hizli koordinasyon ayn sentez akna baglanmaldr.",
            "standards": [
                "Bolgesel gelismeler savunma ve ekonomi etkileri ile birlikte ele alinmaldr.",
                "Uluslararasi sinyaller gecikmeli degil es zamanli risk tablosuna islenmelidir.",
                "Karar alicilar icin alternatif senaryo dili net ve olculebilir tutulmaldr.",
            ],
        },
        "toplumsal_olaylar": {
            "cover_title": "TOPLUMSAL OLAYLAR DURUM DEGERLENDIRMESI",
            "scope": "Toplumsal hareketlilik, kamu duzeni, sahadaki yansmalar ve kriz tetikleyiciler",
            "capacity": "Toplumsal gorunum; saha ivmesi, kurumsal reaksiyon kabiliyeti ve kamu duzeni etkileri uzerinden degerlendirilmistir.",
            "architecture": "Yerel sinyal, teyitli saha notu, merkez alarm hatt ve operasyon mesajlasmas tek aksta tutulmaldr.",
            "standards": [
                "Saha teyidi olmadan alarm seviyeleri nihai karar dili gibi kullanlmamaldr.",
                "Yerel, bolgesel ve ulusal etkiler ayri fakat baglantili izlenmelidir.",
                "Kamu duzeni ve toplumsal alg etkisi ayn rapor omurgasnda birlestirilmelidir.",
            ],
        },
        "genel_chat": {
            "cover_title": "GENEL CHAT GORUSME NOTU",
            "scope": "Serbest soru-cevap, genel bilgi, hizli yorum ve sohbet ozeti",
            "capacity": "Bu mod resmi analiz yerine sohbet odakl dogrudan cevap uretir.",
            "architecture": "Kullanici sorusu ile sohbet cevabi arasndaki akis dogrudan ve sade tutulmustur.",
            "standards": [
                "Sohbet notu ksa ve net tutulmaldr.",
                "Cevaplar dogrudan kullanici mesajna bagli olmaldr.",
                "Genel Chat ciktisi resmi analiz yerine gorusme notu olarak degerlendirilmelidir.",
            ],
        },
        "cross": {
            "cover_title": "CAPRAZ ETKI STRATEJIK SENTEZ RAPORU",
            "scope": "Savunma, ekonomi, enerji, dis politika ve toplumsal olaylarin birlesik etkisi",
            "capacity": "Capraz sentez; alanlar arasi etkilesim, zincirleme risk ve merkez karar basks uzerinden okunmustur.",
            "architecture": "Alan bazl sinyaller tek merkezde sentezlenmeli ve birlesik etki tablosu uretilmelidir.",
            "standards": [
                "Alanlar arasi baglar ayri ayri degil ortak risk matrisi ile sunulmaldr.",
                "Cok alanli analizlerde zaman cizelgesi ve etki siddeti birlikte ele alinmaldr.",
                "Nihai yonetici ozeti tek paragraf degil karar odakl sentez seklinde kurulmaldr.",
            ],
        },
    }
    return profiles.get(domain, profiles["savunma"])


def build_report_package(analysis_id: str, domain: str, situation: str, result: dict) -> dict:
    profile = report_profile_for(domain)
    threat_level = result.get("genel_tehdit_seviyesi") or result.get("tehdit_seviyesi") or "ORTA"
    timestamp = result.get("timestamp") or stamp()
    title = report_title_for(domain)

    phases = [
        {
            "faz": "Faz 1 | Ilk 24 Saat",
            "sure": "0-24 Saat",
            "hedef": "Durumun teyidi, veri birlestirme ve merkez alarm akisinin netlestirilmesi.",
            "sorumlu": ", ".join(result.get("etkilenen_kurumlar")[:2] or DOMAINS[domain]["kurumlar"][:2]),
        },
        {
            "faz": "Faz 2 | 72 Saat",
            "sure": "1-3 Gun",
            "hedef": "Kritik baglantilarin izlenmesi ve kurumlar arasi uygulama hattinin kurulmasi.",
            "sorumlu": ", ".join(result.get("etkilenen_kurumlar")[:3] or DOMAINS[domain]["kurumlar"]),
        },
        {
            "faz": "Faz 3 | Orta Vade",
            "sure": result.get("zaman_cercevesi") or "3-10 Gun",
            "hedef": "Stratejik aksiyonlarin olculmesi ve yeni senaryolara gore duzeltici adimlarin uygulanmasi.",
            "sorumlu": DOMAINS[domain]["kurumlar"][0],
        },
    ]

    risks = []
    for scenario in result.get("senaryolar", [])[:3]:
        if isinstance(scenario, dict):
            risks.append(
                {
                    "baslik": scenario.get("baslik") or "Risk",
                    "aciklama": scenario.get("aciklama") or "Degerlendirme gerekiyor.",
                    "tedbir": scenario.get("aksiyon") or "Yakn izleme onerilir.",
                }
            )
        else:
            risks.append(
                {
                    "baslik": "Risk",
                    "aciklama": str(scenario),
                    "tedbir": "Yakn izleme onerilir.",
                }
            )

    institutions = result.get("etkilenen_kurumlar") or DOMAINS[domain]["kurumlar"]
    kurum_bullets = [
        f"{institution}: alan bazli operasyon akisinda sorumlu veya etkili kurum olarak degerlendirilmelidir."
        for institution in institutions
    ]

    critical_findings = [
        result.get("tehdit_analizi") or "Tehdit analizi hazirlanmistir.",
        result.get("kritik_baglanti") or "Kritik baglantilar izlenmelidir.",
        f"Genel tehdit seviyesi: {threat_level}",
    ]

    recommendations = [result.get("oncelikli_oneri") or "Kurumsal esgudum korunmalidir."]
    for scenario in result.get("senaryolar", [])[:2]:
        if isinstance(scenario, dict):
            recommendations.append(scenario.get("aksiyon") or "Ek aksiyon degerlendirilmelidir.")
        else:
            recommendations.append(str(scenario))

    package = {
        "kapak": {
            "kurum": "BOLD Askeri Teknoloji ve Savunma Sanayi A.S.",
            "birim": "Stratejik Analiz ve Politika Gelistirme Birimi",
            "sistem": "T.C. ANATOLIA-Q",
            "proje_kodu": "QTR-202412",
            "cikti_no": analysis_id,
            "belge_no": f"AQ/{report_code_for(domain)}/{datetime.now().strftime('%Y%m%d')}/{analysis_id[-4:]}",
            "tarih": timestamp,
            "baslik": profile["cover_title"],
            "gizlilik": f"GIZLILIK DERECESI: {threat_level}",
            "kapsam": profile["scope"],
        },
        "yonetici_ozeti": result.get("ozet") or "Yonetici ozeti olusturulamadi.",
        "kritik_bulgular": critical_findings,
        "temel_oneriler": recommendations,
        "tehdit_analizi_bolumu": result.get("tehdit_analizi") or "Tehdit analizi hazirlanmamistir.",
        "mevcut_kapasite": profile["capacity"],
        "onerilen_mimari": profile["architecture"],
        "bolgesel_analiz": situation or result.get("kritik_baglanti") or "Bolgesel analiz notu bulunmamaktadir.",
        "uygulama_plani": phases,
        "kurumsal_sorumluluklar": kurum_bullets,
        "teknik_standartlar": profile["standards"],
        "riskler_ve_tedbirler": risks,
        "sonuc_ve_eylem_cagrisi": result.get("oncelikli_oneri") or "Ilgili kurumlarin koordineli bicimde hareket etmesi onerilir.",
        "rapor_basligi": title,
    }

    if domain == "cross":
        alan_etkileri = result.get("alan_etkileri") or {}
        package["alan_etki_ozeti"] = [
            f"{DOMAINS[key]['display']}: {(alan_etkileri.get(key) or {}).get('etki', 'orta')} etki - {(alan_etkileri.get(key) or {}).get('aciklama', 'Takip edilmelidir.') }"
            for key in ["savunma", "ekonomi", "enerji", "toplumsal_olaylar", "dis_politika"]
        ]

    return package


def escape_html(value: str) -> str:
    return str(value or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def html_bullets(items) -> str:
    return "".join(f"<li>{escape_html(item)}</li>" for item in items if item)


def build_analysis_mail_html(actor: dict, domain: str, situation: str, result: dict) -> str:
    kurumlar = result.get("etkilenen_kurumlar") or []
    scenarios = []
    for item in result.get("senaryolar") or []:
        if isinstance(item, dict):
            scenarios.append(
                f"<li><b>{escape_html(item.get('baslik'))}</b> | {escape_html(item.get('olasilik'))}<br>{escape_html(item.get('aciklama'))}<br><i>{escape_html(item.get('aksiyon'))}</i></li>"
            )
        else:
            scenarios.append(f"<li>{escape_html(item)}</li>")

    return f"""
    <h2>T.C. ANATOLIA-Q Analiz Bildirimi</h2>
    <p><b>Kullanici:</b> {escape_html(actor.get('username', 'bilinmiyor'))}</p>
    <p><b>Rol:</b> {escape_html(actor.get('role', 'bilinmiyor'))}</p>
    <p><b>Alan:</b> {escape_html(DOMAINS.get(domain, {}).get('display', domain))}</p>
    <p><b>Analiz No:</b> {escape_html(result.get('analysis_id'))}</p>
    <p><b>Zaman:</b> {escape_html(result.get('timestamp'))}</p>
    <hr>
    <p><b>Durum Notu:</b><br>{escape_html(situation)}</p>
    <p><b>Yonetici Ozeti:</b><br>{escape_html(result.get('ozet'))}</p>
    <p><b>Tehdit Analizi:</b><br>{escape_html(result.get('tehdit_analizi'))}</p>
    <p><b>Oncelikli Oneri:</b><br>{escape_html(result.get('oncelikli_oneri'))}</p>
    <p><b>Zaman Cercevesi:</b> {escape_html(result.get('zaman_cercevesi'))}</p>
    <p><b>Tehdit Seviyesi:</b> {escape_html(result.get('genel_tehdit_seviyesi') or result.get('tehdit_seviyesi'))}</p>
    <p><b>Kritik Baglanti:</b><br>{escape_html(result.get('kritik_baglanti'))}</p>
    <p><b>Etkilenen Kurumlar:</b></p>
    <ul>{html_bullets(kurumlar)}</ul>
    <p><b>Senaryolar:</b></p>
    <ul>{''.join(scenarios)}</ul>
    """


def ensure_docx_support():
    if not all([Document, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Pt, RGBColor]):
        raise HTTPException(500, "python-docx bagimliligi eksik.")


def add_run(paragraph, text, *, size=11, bold=False, color=None, uppercase=False):
    content = str(text or "")
    if uppercase:
        content = content.upper()
    run = paragraph.add_run(content)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_section_heading(document, title, level=1):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(6)
    add_run(paragraph, title, size=15 if level == 1 else 13, bold=True, color="1A3A5C")
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        add_run(paragraph, item, size=11)


def add_phase_table(document, phases):
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Faz", "Sure", "Hedef", "Sorumlu"]
    for index, label in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "1A3A5C")
        add_run(cell.paragraphs[0], label, size=10, bold=True, color="FFFFFF")

    for phase in phases:
        row = table.add_row().cells
        values = [phase.get("faz"), phase.get("sure"), phase.get("hedef"), phase.get("sorumlu")]
        for index, value in enumerate(values):
            add_run(row[index].paragraphs[0], value, size=10)


def build_docx_report(report: dict, result: dict, domain: str) -> BytesIO:
    ensure_docx_support()
    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(42)
    section.bottom_margin = Pt(42)
    section.left_margin = Pt(52)
    section.right_margin = Pt(52)

    kapak = report.get("kapak") or {}
    title = kapak.get("baslik") or "Analiz Raporu"
    domain_name = DOMAINS.get(domain, {}).get("display") or domain
    summary = report.get("yonetici_ozeti") or result.get("ozet") or ""

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, kapak.get("kurum") or "BOLD Askeri Teknoloji ve Savunma Sanayi A.S.", size=15, bold=True, color="1A3A5C", uppercase=True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    add_run(p, kapak.get("birim") or "Stratejik Analiz ve Politika Gelistirme Birimi", size=10, color="6C7A89")

    p = document.add_paragraph()
    add_run(p, kapak.get("sistem") or "T.C. ANATOLIA-Q", size=15, bold=True, color="1A3A5C")

    p = document.add_paragraph()
    add_run(p, f"Proje Kodu: {kapak.get('proje_kodu') or 'QTR-202412'}", size=10, color="6C7A89")

    p = document.add_paragraph()
    add_run(p, f"Sistem Ciktisi No: {kapak.get('cikti_no') or result.get('analysis_id') or '--'}", size=10, color="6C7A89")

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_run(p, title, size=19, bold=True, color="1A3A5C", uppercase=True)

    p = document.add_paragraph()
    add_run(p, domain_name, size=15, bold=True, color="C0392B", uppercase=True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    add_run(p, kapak.get("gizlilik") or "GIZLILIK DERECESI: GIZLI", size=14, bold=True, color="C0392B", uppercase=True)

    if summary:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        add_run(p, summary, size=11)

    meta_table = document.add_table(rows=0, cols=2)
    meta_table.style = "Table Grid"
    meta_items = [
        ("Belge No", kapak.get("belge_no") or result.get("analysis_id") or "--"),
        ("Tarih", kapak.get("tarih") or result.get("timestamp") or "--"),
        ("Hazirlayan", kapak.get("kurum") or ""),
        ("Kapsam", kapak.get("kapsam") or domain_name),
        ("Siniflandirma", kapak.get("gizlilik") or "GIZLI"),
        ("Oncelik", result.get("tehdit_seviyesi") or result.get("genel_tehdit_seviyesi") or "--"),
    ]
    for label, value in meta_items:
        row = meta_table.add_row().cells
        row[0].text = ""
        row[1].text = ""
        set_cell_shading(row[0], "F6F8FB")
        add_run(row[0].paragraphs[0], label, size=10, bold=True, color="1A3A5C")
        add_run(row[1].paragraphs[0], value or "-", size=10)

    add_section_heading(document, "Yonetici Ozeti")
    document.add_paragraph(report.get("yonetici_ozeti") or result.get("ozet") or "")

    add_section_heading(document, "Kritik Bulgular")
    add_bullets(document, report.get("kritik_bulgular") or [])

    add_section_heading(document, "Temel Oneriler")
    add_bullets(document, report.get("temel_oneriler") or [])

    add_section_heading(document, "1. Tehdit Analizi")
    document.add_paragraph(report.get("tehdit_analizi_bolumu") or result.get("tehdit_analizi") or "")

    add_section_heading(document, "2. Mevcut Kapasite Degerlendirmesi")
    document.add_paragraph(report.get("mevcut_kapasite") or "")

    add_section_heading(document, "3. Onerilen Tespit Mimarisi")
    document.add_paragraph(report.get("onerilen_mimari") or result.get("kritik_baglanti") or "")

    add_section_heading(document, "4. Bolge / Alan Bazli Degerlendirme")
    document.add_paragraph(report.get("bolgesel_analiz") or result.get("ozet") or "")

    add_section_heading(document, "5. Uygulama Plani ve Zaman Cizelgesi")
    add_phase_table(document, report.get("uygulama_plani") or [])

    add_section_heading(document, "6. Kurumsal Yapi ve Sorumluluklar")
    add_bullets(document, report.get("kurumsal_sorumluluklar") or [])

    add_section_heading(document, "7. Teknik Standartlar ve Minimum Gereksinimler")
    add_bullets(document, report.get("teknik_standartlar") or [])

    add_section_heading(document, "8. Riskler ve Azaltici Tedbirler")
    for item in report.get("riskler_ve_tedbirler") or []:
        paragraph = document.add_paragraph(style="List Bullet")
        add_run(paragraph, f"{item.get('baslik') or 'Risk'}: ", size=11, bold=True, color="1A3A5C")
        add_run(paragraph, item.get("aciklama") or "", size=11)
        if item.get("tedbir"):
            add_run(paragraph, f" Tedbir: {item.get('tedbir')}", size=11, color="C0392B")

    add_section_heading(document, "9. Sonuc ve Eylem Cagrisi")
    document.add_paragraph(report.get("sonuc_ve_eylem_cagrisi") or result.get("oncelikli_oneri") or "")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def parse_provider_json(raw_text: str) -> dict:
    text = extract_json_text(raw_text)
    data = json.loads(text)
    if not isinstance(data, dict):
        raise ValueError("Saglayc JSON nesnesi dondurmedi.")
    return data


def call_openai_like(provider_name: str, api_key: str, model: str, domain: str, situation: str, chat_name: str = "", base_url: str | None = None):
    if OpenAI is None:
        raise RuntimeError(f"{provider_name} istemcisi kullanlamyor.")
    if not model:
        raise RuntimeError(f"{provider_name} model bilgisi eksik.")

    kwargs = {"api_key": api_key} if api_key else {}
    if base_url:
        kwargs["base_url"] = base_url
    client = OpenAI(**kwargs)

    response = client.chat.completions.create(
        model=model,
        temperature=0.35,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": build_system_prompt(domain)},
            {"role": "user", "content": build_user_prompt(domain, situation, chat_name)},
        ],
    )
    text = response.choices[0].message.content or "{}"
    parsed = parse_provider_json(text)
    return normalize_analysis(domain, parsed, chat_name), {"provider": provider_name, "model": model}


def call_anthropic(domain: str, situation: str, chat_name: str = ""):
    if Anthropic is None or not ANTHROPIC_API_KEY:
        raise RuntimeError("Anthropic yaplandrlmams.")

    client = Anthropic(api_key=ANTHROPIC_API_KEY)
    response = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=1800,
        temperature=0.35,
        system=build_system_prompt(domain),
        messages=[
            {
                "role": "user",
                "content": build_user_prompt(domain, situation, chat_name) + "\n\nYalnzca gecerli JSON dondur.",
            }
        ],
    )
    parts = [block.text for block in response.content if getattr(block, "type", "") == "text"]
    parsed = parse_provider_json("\n".join(parts))
    return normalize_analysis(domain, parsed, chat_name), {"provider": "anthropic", "model": ANTHROPIC_MODEL}


def analyze_with_router(domain: str, situation: str, chat_name: str = ""):
    route = PROVIDER_ROUTE.get(domain, ["anthropic", "openai", "grok", "llama", "fallback"])
    errors = []
    availability = provider_status()

    for provider in route:
        if not availability.get(provider):
            errors.append(f"{provider}:not-configured")
            continue

        try:
            if provider == "openai":
                result, meta = call_openai_like("openai", OPENAI_API_KEY, OPENAI_MODEL, domain, situation, chat_name)
            elif provider == "anthropic":
                result, meta = call_anthropic(domain, situation, chat_name)
            elif provider == "grok":
                result, meta = call_openai_like("grok", XAI_API_KEY, XAI_MODEL, domain, situation, chat_name, base_url=XAI_BASE_URL)
            elif provider == "llama":
                result, meta = call_openai_like("llama", LLAMA_API_KEY, LLAMA_MODEL, domain, situation, chat_name, base_url=LLAMA_BASE_URL)
            else:
                result = fallback(domain, situation, chat_name)
                meta = {"provider": "fallback", "model": "fallback-core"}

            result["_meta"] = {
                "provider": meta["provider"],
                "model": meta["model"],
                "route": route,
                "attempt_errors": errors,
                "fallback_mode": meta["provider"] == "fallback",
            }
            return result
        except Exception as error:
            errors.append(f"{provider}:{error.__class__.__name__}")
            if is_soft_provider_failure(error) and provider in {"openai", "anthropic", "grok", "llama"}:
                cool_down_provider(provider)
            if not is_soft_provider_failure(error) and provider == "fallback":
                raise
            continue

    result = fallback(domain, situation, chat_name)
    result["_meta"] = {
        "provider": "fallback",
        "model": "fallback-core",
        "route": route,
        "attempt_errors": errors,
        "fallback_mode": True,
    }
    return result


def save_analysis(domain, situation, result, actor=None):
    analysis_id = "AQ-" + uuid.uuid4().hex[:6].upper()
    created = stamp()
    payload = normalize_payload(dict(result))
    meta = dict(payload.get("_meta") or {})
    actor = actor or {}
    payload.update(
        {
            "analysis_id": analysis_id,
            "timestamp": created,
            "time": created,
            "created_at": created,
            "fallback_mode": bool(meta.get("fallback_mode")),
            "provider": meta.get("provider", "fallback"),
            "model": meta.get("model", "fallback-core"),
            "risk_analizi": payload.get("tehdit_analizi", ""),
            "actor_username": actor.get("username", "bilinmiyor"),
            "actor_role": actor.get("role", "bilinmiyor"),
        }
    )
    payload["senaryo_analizi"] = [
        f"{item['baslik']} | Olaslk: {item['olasilik']} | {item['aciklama']} | Aksiyon: {item['aksiyon']}"
        for item in payload.get("senaryolar", [])
        if isinstance(item, dict)
    ]
    payload["report_package"] = build_report_package(analysis_id, domain, situation, payload)
    analysis_store[analysis_id] = {
        "id": analysis_id,
        "domain": domain,
        "timestamp": created,
        "user": payload["actor_username"],
        "role": payload["actor_role"],
        "result": payload,
    }
    safe_send_mail(
        f"T.C. ANATOLIA-Q Analiz | {DOMAINS.get(domain, {}).get('display', domain)} | {analysis_id}",
        build_analysis_mail_html(actor, domain, situation, payload),
    )
    return payload


def patch_frontend(html: str) -> str:
    inject = (
        f'<script src="/chat-hotfix.js?v={APP_VERSION}"></script>'
        f'<script src="/report-hotfix.js?v={APP_VERSION}"></script>'
        f'<script src="/ui-tidy-hotfix.js?v={APP_VERSION}"></script>'
    )
    marker = "</body>"
    index = html.rfind(marker)
    if index == -1:
        return html + inject
    return html[:index] + inject + html[index:]


@app.get("/")
async def root():
    path = os.path.join(os.path.dirname(__file__), "index.html")
    if not os.path.exists(path):
        return HTMLResponse("<h1>T.C. ANATOLIA-Q</h1>")
    with open(path, "r", encoding="utf-8") as handle:
        return HTMLResponse(patch_frontend(handle.read()))


@app.get("/chat-hotfix.js")
async def chat_hotfix():
    path = os.path.join(os.path.dirname(__file__), "chat_hotfix.js")
    if not os.path.exists(path):
        return Response("// chat hotfix missing", media_type="application/javascript")
    with open(path, "r", encoding="utf-8") as handle:
        return Response(handle.read(), media_type="application/javascript")


@app.get("/ui-tidy-hotfix.js")
async def ui_tidy_hotfix():
    path = os.path.join(os.path.dirname(__file__), "ui-tidy-hotfix.js")
    if not os.path.exists(path):
        return Response("// layout hotfix missing", media_type="application/javascript")
    with open(path, "r", encoding="utf-8") as handle:
        return Response(handle.read(), media_type="application/javascript")


@app.get("/report-hotfix.js")
async def report_hotfix():
    path = os.path.join(os.path.dirname(__file__), "report-hotfix.js")
    if not os.path.exists(path):
        return Response("// report hotfix missing", media_type="application/javascript")
    with open(path, "r", encoding="utf-8") as handle:
        return Response(handle.read(), media_type="application/javascript")


@app.get("/health")
async def health():
    return {
        "status": "online",
        "system": "T.C. ANATOLIA-Q",
        "version": APP_VERSION,
        "provider": "multi-router",
        "providers": provider_status(),
        "routes": PROVIDER_ROUTE,
    }


@app.post("/api/login")
async def login(data: dict):
    username = str(data.get("username", "")).strip()
    password = str(data.get("password", ""))
    if username not in USERS or password != PASSWORD:
        raise HTTPException(401, "Kullanc kodu veya sifre hatal.")

    code = "".join(random.choices(string.digits, k=6))
    pending_codes[username] = {"code": code, "expires": datetime.now() + timedelta(minutes=10)}
    send_mail(
        "T.C. ANATOLIA-Q Dogrulama Kodu",
        f"<p>Kullanc: <b>{username}</b></p><p>Kod: <b>{code}</b></p><p>Saat: {stamp()}</p>",
    )
    return {
        "message": "Dogrulama kodu gonderildi.",
        "email_hint": PRIMARY_EMAIL[:3] + "***@" + PRIMARY_EMAIL.split("@")[-1],
    }


@app.post("/api/verify")
async def verify(data: dict):
    username = str(data.get("username", "")).strip()
    code = str(data.get("code", "")).strip()
    pending = pending_codes.get(username)

    if username not in USERS:
        raise HTTPException(401, "Gecersiz kullanc.")
    if not pending:
        raise HTTPException(401, "Kod bulunamad. Tekrar giris yapn.")
    if datetime.now() > pending["expires"]:
        del pending_codes[username]
        raise HTTPException(401, "Kodun suresi doldu.")
    if pending["code"] != code:
        raise HTTPException(401, "Hatal dogrulama kodu.")

    del pending_codes[username]
    token = f"aq_{username}_{uuid.uuid4().hex[:16]}"
    active_sessions[token] = {"username": username, **USERS[username]}
    return {
        "token": token,
        "name": USERS[username]["name"],
        "role": USERS[username]["role"],
        "username": username,
        "user": username,
    }


@app.post("/api/contact-center")
async def contact_center(request: Request, req: dict):
    session = require_session(request, req)
    note = str(req.get("note", "")).strip()
    send_mail(
        "T.C. ANATOLIA-Q Merkez Iletisim Talebi",
        f"<p>Kullanc: <b>{session['username']}</b></p><p>Rol: <b>{session['role']}</b></p><p>Not: {note or '-'}</p><p>Saat: {stamp()}</p>",
    )
    return {"message": "Merkez iletisim talebiniz gonderildi."}


@app.get("/api/alerts")
async def alerts():
    return {"items": list(reversed(alerts_store))}


@app.post("/api/alerts")
async def create_alert(request: Request, req: dict):
    session = require_session(request, req)
    region = str(req.get("region", "")).strip() or "Genel"
    title = str(req.get("title", "")).strip() or "Bolgesel alarm"
    detail = str(req.get("detail", "")).strip() or "Merkez inceleme bekliyor."
    priority = str(req.get("priority", "")).strip().upper() or "ORTA"

    item = {
        "id": "ALM-" + uuid.uuid4().hex[:6].upper(),
        "region": region[:32],
        "title": title[:80],
        "detail": detail[:400],
        "priority": priority[:16],
        "user": session["username"],
        "role": session["role"],
        "timestamp": stamp(),
    }
    alerts_store.append(item)
    trim_store(alerts_store)

    try:
        send_mail(
            f"T.C. ANATOLIA-Q Alarm | {item['region']}",
            f"<p>Kullanc: <b>{session['username']}</b></p><p>Rol: <b>{session['role']}</b></p><p>Bolge: <b>{item['region']}</b></p><p>Baslk: <b>{item['title']}</b></p><p>Detay: {item['detail']}</p><p>Oncelik: <b>{item['priority']}</b></p><p>Saat: {item['timestamp']}</p>",
        )
    except HTTPException:
        pass

    ops_feed_store.append(
        {
            "id": "MSG-" + uuid.uuid4().hex[:6].upper(),
            "channel": "alarm",
            "priority": item["priority"],
            "message": f"{item['region']} bolgesi icin alarm gecti: {item['title']}",
            "user": session["username"],
            "role": session["role"],
            "timestamp": item["timestamp"],
        }
    )
    trim_store(ops_feed_store)
    return {"message": "Alarm tum kullanclar icin kaydedildi.", "item": item}


@app.get("/api/ops-feed")
async def ops_feed():
    return {"items": list(reversed(ops_feed_store))}


@app.post("/api/ops-feed")
async def push_ops_feed(request: Request, req: dict):
    session = require_session(request, req)
    message = str(req.get("message", "")).strip()
    channel = str(req.get("channel", "")).strip() or "ops"
    priority = str(req.get("priority", "")).strip().upper() or "BILGI"

    if not message:
        raise HTTPException(400, "Mesaj bos olamaz.")

    item = {
        "id": "MSG-" + uuid.uuid4().hex[:6].upper(),
        "channel": channel[:24],
        "priority": priority[:16],
        "message": message[:600],
        "user": session["username"],
        "role": session["role"],
        "timestamp": stamp(),
    }
    ops_feed_store.append(item)
    trim_store(ops_feed_store, limit=120)

    if channel == "emergency":
        try:
            send_mail(
                "T.C. ANATOLIA-Q Acil Alarm",
                f"<p>Kullanc: <b>{session['username']}</b></p><p>Rol: <b>{session['role']}</b></p><p>Mesaj: {item['message']}</p><p>Oncelik: <b>{item['priority']}</b></p><p>Saat: {item['timestamp']}</p>",
            )
        except HTTPException:
            pass

    return {"message": "Operasyon aks guncellendi.", "item": item}


@app.post("/api/analyze")
async def analyze(request: Request, req: dict):
    domain = str(req.get("domain", "")).strip()
    situation = str(req.get("situation", "")).strip()
    chat_name = str(req.get("chat_name", "")).strip()

    if domain not in DOMAINS:
        raise HTTPException(400, "Gecersiz alan.")
    if not situation:
        raise HTTPException(400, "Durum bildirimi bos.")

    actor = active_sessions.get(token_from(request, req), {})
    return save_analysis(domain, situation, analyze_with_router(domain, situation, chat_name), actor=actor)


@app.get("/api/history")
async def history():
    items = sorted(analysis_store.values(), key=lambda item: item["timestamp"], reverse=True)
    return [
        {
            "id": item["id"],
            "domain": item["domain"],
            "dom": item["domain"],
            "timestamp": item["result"]["timestamp"],
            "time": item["result"]["timestamp"],
            "summary": str(item["result"].get("ozet", ""))[:130],
            "ozet": str(item["result"].get("ozet", ""))[:130],
            "fallback_mode": bool(item["result"].get("fallback_mode")),
            "provider": item["result"].get("provider", "fallback"),
            "model": item["result"].get("model", "fallback-core"),
            "user": item.get("user") or item["result"].get("actor_username", "bilinmiyor"),
            "role": item.get("role") or item["result"].get("actor_role", "bilinmiyor"),
            "report_docx_url": f"/api/report/{item['id']}/docx",
            "result": item["result"],
        }
        for item in items[:20]
    ]


@app.get("/api/report/{analysis_id}/docx")
async def report_docx(analysis_id: str, request: Request):
    item = analysis_store.get(analysis_id)
    if not item:
        raise HTTPException(404, "Rapor bulunamadi.")

    result = item["result"]
    report = result.get("report_package") or build_report_package(analysis_id, item["domain"], "", result)
    filename = f"TC_ANATOLIA_Q_{analysis_id}_{item['domain']}.docx"
    buffer = build_docx_report(report, result, item["domain"])
    buffer.seek(0)
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
